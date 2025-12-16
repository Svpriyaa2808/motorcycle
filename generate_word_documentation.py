#!/usr/bin/env python3
"""
Generate comprehensive Word documentation for the Motorcycle Repair Shops project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)  # Blue
        heading.runs[0].font.size = Pt(24)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)  # Lighter blue
        heading.runs[0].font.size = Pt(18)
    return heading

def add_paragraph(doc, text, bold=False):
    """Add a paragraph with custom formatting"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(12)
    if bold:
        para.runs[0].bold = True
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.space_after = Pt(6)
    return para

def create_word_documentation():
    """Generate comprehensive Word documentation"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('🏍️ Motorcycle Repair Shops', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('European Directory', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    project_doc = doc.add_paragraph('Project Documentation')
    project_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_doc.runs[0].font.size = Pt(16)
    project_doc.runs[0].font.color.rgb = RGBColor(55, 65, 81)

    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Technology Stack',
        '3. Development Timeline',
        '4. Key Features',
        '5. Architecture & Components',
        '6. Data Management',
        '7. Setup & Installation',
        '8. Project Structure',
        '9. Future Enhancements',
    ]
    for item in toc_items:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # 1. Project Overview
    add_heading(doc, '1. Project Overview', level=1)

    add_paragraph(doc,
        'The Motorcycle Repair Shops European Directory is a modern, completely FREE web application '
        'designed to help motorcycle enthusiasts and riders find reliable repair shops across 27 European '
        'countries. Built with cutting-edge web technologies, this application provides an intuitive '
        'interface for discovering, searching, and locating motorcycle service centers throughout Europe.'
    )

    add_heading(doc, 'Key Highlights:', level=2)
    highlights = [
        'Covers 27 European Union countries',
        '3000+ motorcycle repair shops',
        '100% FREE - No API costs, no billing required',
        'Interactive maps powered by OpenStreetMap',
        'Fully responsive design for all devices',
        'Real-time search and filtering capabilities',
        'Modern, gradient-based UI design',
    ]
    for highlight in highlights:
        add_bullet_point(doc, highlight)

    doc.add_page_break()

    # 2. Technology Stack
    add_heading(doc, '2. Technology Stack', level=1)

    add_heading(doc, 'Frontend Technologies:', level=2)
    add_paragraph(doc, 'Next.js 16:', bold=True)
    add_paragraph(doc,
        "The latest version of React's premier framework, providing server-side rendering, "
        "static site generation, and optimized performance out of the box."
    )

    add_paragraph(doc, 'React 19:', bold=True)
    add_paragraph(doc,
        'The newest version of React, offering improved performance, enhanced hooks, '
        'and better developer experience with concurrent features.'
    )

    add_paragraph(doc, 'TypeScript:', bold=True)
    add_paragraph(doc,
        'Provides type safety throughout the application, reducing bugs and improving code maintainability.'
    )

    add_paragraph(doc, 'Tailwind CSS 4:', bold=True)
    add_paragraph(doc,
        'Utility-first CSS framework enabling rapid UI development with consistent design patterns '
        'and responsive layouts.'
    )

    add_heading(doc, 'Mapping Solution:', level=2)
    add_paragraph(doc, 'Leaflet + OpenStreetMap:', bold=True)
    add_paragraph(doc,
        'A completely FREE mapping solution that requires no API keys, no billing setup, and has no usage limits. '
        'Leaflet is a lightweight JavaScript library for interactive maps, while OpenStreetMap provides '
        'community-driven, open-source map data.'
    )

    add_heading(doc, 'Data Management:', level=2)
    add_paragraph(doc, 'CSV File Storage:', bold=True)
    add_paragraph(doc,
        'Shop data is stored in CSV format for easy management and portability. The application includes '
        'a custom CSV parser that handles quoted fields and complex data structures.'
    )

    add_paragraph(doc, 'Supabase (Optional):', bold=True)
    add_paragraph(doc,
        'PostgreSQL-based backend service for advanced features and real-time capabilities. '
        'The free tier is sufficient for most use cases.'
    )

    doc.add_page_break()

    # 3. Development Timeline
    add_heading(doc, '3. Development Timeline', level=1)

    add_paragraph(doc,
        'The project was developed through several iterative phases, each adding significant '
        'functionality and improvements:'
    )

    phases = [
        ('Phase 1 (Week 1) - Initial Setup', [
            'Created Next.js project',
            'Basic project structure',
            'Initial commit',
        ]),
        ('Phase 2 (Week 2-3) - Core Features', [
            'Country filtering system',
            'Google Maps integration (later replaced)',
            'Modern UI with gradients',
            'Responsive design implementation',
        ]),
        ('Phase 3 (Week 4) - Configuration & Docs', [
            'Comprehensive setup guides',
            'API key configuration',
            'Setup checker script',
            'Quick start documentation',
        ]),
        ('Phase 4 (Week 5) - Maps Migration', [
            'Replaced Google Maps with OpenStreetMap',
            'Integrated Leaflet library',
            'Removed API key requirements',
            'Made app 100% FREE to use',
        ]),
        ('Phase 5 (Week 6-7) - Data Enhancement', [
            'Added CSV file support',
            'Built custom CSV parser',
            'Migrated to CSV data storage',
            'Data setup documentation',
        ]),
        ('Phase 6 (Recent) - Optimization', [
            'Fixed shop list display issues',
            'Optimized for 3000+ shops',
            'Performance improvements',
            'Bug fixes and refinements',
        ]),
    ]

    for phase_title, achievements in phases:
        add_heading(doc, phase_title, level=2)
        for achievement in achievements:
            add_bullet_point(doc, achievement)

    doc.add_page_break()

    # 4. Key Features
    add_heading(doc, '4. Key Features', level=1)

    add_heading(doc, '4.1 Interactive Maps', level=2)
    add_paragraph(doc,
        'The application features a fully interactive map powered by Leaflet and OpenStreetMap. Users can:'
    )
    map_features = [
        'View all motorcycle shops as custom-styled markers on the map',
        'Zoom and pan to explore different regions',
        'Click on markers to see detailed shop information in popups',
        'Automatic map centering based on filtered results',
        'Smart bounds adjustment when filtering by country',
        'Custom motorcycle-themed markers with gradient styling',
    ]
    for feature in map_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The map automatically adjusts its view based on the selected shops, ensuring users always '
        'see the most relevant information without manual navigation.'
    )

    add_heading(doc, '4.2 Country Filtering System', level=2)
    add_paragraph(doc,
        'A sophisticated country filtering system allows users to narrow down their search:'
    )
    country_features = [
        'Desktop view displays an elegant grid of all 27 EU countries',
        'Mobile view uses a dropdown selector for space efficiency',
        'Each country displays its flag emoji and name',
        'Real-time shop count updates when countries are selected',
        '"All Countries" option to reset the filter',
        'Smooth animations and hover effects for better UX',
    ]
    for feature in country_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The system covers: Austria, Belgium, Bulgaria, Croatia, Cyprus, Czech Republic, Denmark, '
        'Estonia, Finland, France, Germany, Greece, Hungary, Ireland, Italy, Latvia, Lithuania, '
        'Luxembourg, Malta, Netherlands, Poland, Portugal, Romania, Slovakia, Slovenia, Spain, and Sweden.'
    )

    add_heading(doc, '4.3 Smart Search Functionality', level=2)
    add_paragraph(doc,
        'The application includes a powerful search system that enables users to quickly find specific shops:'
    )
    search_features = [
        'Search by shop name, city, street address, or country code',
        'Real-time filtering as users type',
        'Case-insensitive search for better usability',
        'Search results counter showing matching shops',
        'Works in combination with country filtering',
        'Optimized for performance even with 3000+ shops',
    ]
    for feature in search_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        "The search functionality uses React's useMemo hook to optimize performance, ensuring smooth "
        'filtering even with large datasets.'
    )

    add_heading(doc, '4.4 Dual View Modes', level=2)
    add_paragraph(doc, 'Users can switch between two distinct viewing modes:')
    add_paragraph(doc, 'Map View:', bold=True)
    add_paragraph(doc,
        'Displays all shops on an interactive map with custom markers. Perfect for geographical '
        'exploration and finding shops in specific areas.'
    )
    add_paragraph(doc, 'List View:', bold=True)
    add_paragraph(doc, 'Shows shops in a beautiful card-based grid layout. Each card includes:')
    list_features = [
        'Shop name with gradient header',
        'Complete address information',
        'Contact details (phone, email, website)',
        'Country flag and name',
        'Direct link to Google Maps',
        'Hover effects and smooth transitions',
    ]
    for feature in list_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The toggle between views is seamless, maintaining the current filter and search state.'
    )

    add_heading(doc, '4.5 Responsive Design', level=2)
    add_paragraph(doc, 'The application is fully responsive and optimized for all device sizes:')

    add_paragraph(doc, 'Desktop (1920px+):', bold=True)
    add_paragraph(doc,
        'Full-featured interface with grid-based country selector, large map view, and 3-column shop cards.'
    )

    add_paragraph(doc, 'Laptop (1024px - 1919px):', bold=True)
    add_paragraph(doc,
        'Optimized layout with 2-column shop cards and adjusted spacing for comfortable viewing.'
    )

    add_paragraph(doc, 'Tablet (768px - 1023px):', bold=True)
    add_paragraph(doc,
        'Adaptive layout with 2-column cards, touch-optimized controls, and appropriate font sizes.'
    )

    add_paragraph(doc, 'Mobile (320px - 767px):', bold=True)
    add_paragraph(doc,
        'Single-column layout, dropdown country selector, stacked controls, and touch-friendly interface elements.'
    )

    add_paragraph(doc,
        'All interactions are optimized for touch and mouse input, ensuring a great experience '
        'regardless of device type.'
    )

    doc.add_page_break()

    # 5. Architecture & Components
    add_heading(doc, '5. Architecture & Components', level=1)

    add_paragraph(doc,
        'The application follows a modern React architecture with clear separation of concerns:'
    )

    add_heading(doc, '5.1 MotorcycleShops Component', level=2)
    add_paragraph(doc, 'The main component that orchestrates the entire application:')
    comp1_features = [
        'Manages application state (shops, filters, search, view mode)',
        'Fetches data from CSV file on component mount',
        'Implements useMemo for optimized filtering',
        'Renders header with statistics',
        'Coordinates child components',
        'Handles view mode switching',
    ]
    for feature in comp1_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'This component serves as the container for all other components and manages the data flow '
        'throughout the application.'
    )

    add_heading(doc, '5.2 CountrySelector Component', level=2)
    add_paragraph(doc, 'A smart component that provides country filtering functionality:')
    comp2_features = [
        'Displays all 27 EU countries with flags',
        'Responsive grid layout (desktop) / dropdown (mobile)',
        'Highlights selected country',
        'Passes selection back to parent component',
        'Smooth animations and hover effects',
    ]
    for feature in comp2_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The component automatically adapts its layout based on screen size, providing the best '
        'user experience for each device type.'
    )

    add_heading(doc, '5.3 ShopMap Component', level=2)
    add_paragraph(doc, 'An advanced mapping component built with React-Leaflet:')
    comp3_features = [
        'Dynamically imported to prevent SSR issues',
        'Custom map markers with motorcycle icons',
        'Automatic bounds adjustment based on shops',
        'Interactive popups with shop information',
        'Links to external mapping services',
        'Optimized rendering for large datasets',
    ]
    for feature in comp3_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The component includes a MapBoundsHandler sub-component that automatically adjusts the map '
        'view when the shop list changes, ensuring users always see relevant data.'
    )

    add_heading(doc, '5.4 CSV Parser Utility', level=2)
    add_paragraph(doc, 'A custom-built CSV parsing utility that handles data loading:')
    comp4_features = [
        'Parses CSV files with proper quote handling',
        'Transforms flat CSV data into structured objects',
        'Maps country names to ISO codes',
        'Handles missing or malformed data gracefully',
        'Supports various CSV field formats',
    ]
    for feature in comp4_features:
        add_bullet_point(doc, feature)

    add_paragraph(doc,
        'The parser is designed to handle the complexities of real-world CSV data, including quoted '
        'fields, commas within values, and varying data quality.'
    )

    doc.add_page_break()

    # 6. Data Management
    add_heading(doc, '6. Data Management', level=1)

    add_paragraph(doc,
        'The application uses a CSV-based data management system for simplicity and portability:'
    )

    add_heading(doc, 'Data Source:', level=2)
    add_paragraph(doc,
        "All motorcycle shop data is sourced from OpenStreetMap (OSM), the world's largest "
        'collaborative mapping project. The data includes shops tagged as:'
    )
    data_tags = [
        'shop=motorcycle - Dedicated motorcycle shops',
        'craft=motorcycle - Motorcycle craft and repair services',
        'amenity=car_repair with motorcycle=yes - Multi-service repair shops',
    ]
    for tag in data_tags:
        add_bullet_point(doc, tag)

    add_paragraph(doc,
        'Data is fetched using the Overpass API, which allows querying OSM data with complex filters.'
    )

    add_heading(doc, 'Data Structure:', level=2)
    add_paragraph(doc, 'Each shop record contains the following information:')
    data_fields = [
        'Unique ID',
        'Shop name',
        'Geographic coordinates (latitude, longitude)',
        'Complete address (street, house number, postal code, city)',
        'Country code (ISO 3166-1 alpha-2)',
        'Contact information (phone, email, website)',
        'Additional tags from OpenStreetMap',
    ]
    for field in data_fields:
        add_bullet_point(doc, field)

    add_paragraph(doc,
        'The data is stored in a CSV file located at /public/data/eu_motorcycle_repairs.csv, making '
        'it easy to update and maintain.'
    )

    add_heading(doc, 'Data Updates:', level=2)
    add_paragraph(doc,
        'The project includes a Python script (scripts/fetch_osm_data.py) that can fetch fresh data '
        'from OpenStreetMap. This allows the database to be updated with new shops and changes to '
        'existing entries. The script can be run with: npm run fetch:data'
    )

    add_paragraph(doc, 'This ensures the application always has access to the most current information.')

    doc.add_page_break()

    # 7. Setup & Installation
    add_heading(doc, '7. Setup & Installation', level=1)

    add_paragraph(doc,
        'Setting up the project is straightforward and requires minimal configuration:'
    )

    add_heading(doc, 'Prerequisites:', level=2)
    prereqs = [
        'Node.js 18 or higher',
        'npm or yarn package manager',
        'Git for version control',
        '(Optional) Supabase account for advanced features',
    ]
    for prereq in prereqs:
        add_bullet_point(doc, prereq)

    add_paragraph(doc,
        'Important: No API keys are required for maps! The application uses OpenStreetMap, '
        'which is completely free and requires no registration.'
    )

    add_heading(doc, 'Installation Steps:', level=2)
    steps = [
        'Clone the repository: git clone <repository-url>',
        'Navigate to directory: cd motorcycle',
        'Install dependencies: npm install',
        '(Optional) Set up environment variables for Supabase',
        'Start the development server: npm run dev',
        'Open your browser to http://localhost:3000',
    ]
    for i, step in enumerate(steps, 1):
        add_bullet_point(doc, f'{i}. {step}')

    add_heading(doc, 'Available Scripts:', level=2)
    scripts = [
        'npm run dev - Start development server with hot reload',
        'npm run build - Create optimized production build',
        'npm start - Start production server',
        'npm run fetch:data - Fetch latest data from OpenStreetMap',
    ]
    for script in scripts:
        add_bullet_point(doc, script)

    doc.add_page_break()

    # 8. Project Structure
    add_heading(doc, '8. Project Structure', level=1)

    structure = doc.add_paragraph(
        'motorcycle/\n'
        '├── src/\n'
        '│   ├── app/                      # Next.js App Router\n'
        '│   │   ├── layout.tsx           # Root layout component\n'
        '│   │   ├── page.tsx             # Home page\n'
        '│   │   ├── globals.css          # Global styles\n'
        '│   │   └── favicon.ico          # Site favicon\n'
        '│   ├── components/\n'
        '│   │   ├── CountrySelector/     # Country filtering component\n'
        '│   │   ├── MotorcycleShops/     # Main shops component\n'
        '│   │   └── ShopMap/             # Interactive map component\n'
        '│   ├── data/\n'
        '│   │   └── countries.ts         # EU countries configuration\n'
        '│   └── utils/\n'
        '│       └── csvParser.ts         # CSV parsing utility\n'
        '├── public/\n'
        '│   └── data/\n'
        '│       └── eu_motorcycle_repairs.csv  # Shop data (3000+ entries)\n'
        '├── scripts/\n'
        '│   └── fetch_osm_data.py        # OSM data fetching script\n'
        '├── supabase/\n'
        '│   └── supabaseClient.js        # Supabase configuration\n'
        '├── package.json                  # Dependencies and scripts\n'
        '├── tsconfig.json                 # TypeScript configuration\n'
        '├── tailwind.config.js            # Tailwind CSS configuration\n'
        '├── next.config.ts                # Next.js configuration\n'
        '└── README.md                     # Project documentation'
    )
    structure.style = 'No Spacing'
    for run in structure.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    add_paragraph(doc,
        'The project follows Next.js 13+ App Router conventions with a clear separation between '
        'presentation components, utilities, and data. This structure makes the codebase easy to '
        'navigate and maintain.'
    )

    doc.add_page_break()

    # 9. Future Enhancements
    add_heading(doc, '9. Future Enhancements', level=1)

    add_paragraph(doc,
        'The project has a strong foundation and several potential enhancements could further '
        'improve its value:'
    )

    enhancements = [
        ('User Reviews', 'Allow users to rate and review shops for community-driven quality insights'),
        ('Advanced Filters', 'Filter by services, certifications, and ratings for more precise discovery'),
        ('Route Planning', 'Integration with routing services to help plan motorcycle trips'),
        ('Favorites System', 'Save preferred shops to a personal list for quick access'),
        ('Mobile App', 'Native iOS/Android applications for better mobile experience'),
        ('Offline Support', 'Progressive Web App features for access without internet'),
        ('Multilingual', 'Support for multiple European languages for wider accessibility'),
        ('Business Portal', 'Allow shops to claim and update listings for more accurate information'),
    ]

    for title, description in enhancements:
        add_paragraph(doc, f'{title}:', bold=True)
        add_paragraph(doc, description)

    doc.add_page_break()

    # Conclusion
    add_heading(doc, 'Conclusion', level=1)

    add_paragraph(doc,
        'The Motorcycle Repair Shops European Directory represents a modern approach to solving a '
        'real-world problem: helping motorcycle riders find reliable service centers across Europe. '
        'By leveraging cutting-edge web technologies and open-source mapping solutions, the project '
        'delivers a professional, feature-rich application that is completely free to use and deploy.'
    )

    add_paragraph(doc,
        'The development journey showcased in this document demonstrates the evolution from initial '
        'concept through multiple iterations, each adding value and improving the user experience. '
        'The decision to migrate from Google Maps to OpenStreetMap particularly highlights the '
        "project's commitment to accessibility and eliminating barriers to entry."
    )

    add_paragraph(doc,
        'With 3000+ shops across 27 countries, responsive design, intelligent search and filtering, '
        'and an intuitive interface, this application provides real value to the European motorcycle '
        'community while serving as an excellent example of modern web development practices.'
    )

    doc.add_paragraph()
    doc.add_paragraph('═' * 80)
    doc.add_paragraph()

    footer_info = doc.add_paragraph()
    footer_info.add_run('Project: ').bold = True
    footer_info.add_run('Motorcycle Repair Shops - European Directory\n')
    footer_info.add_run('Documentation Generated: ').bold = True
    footer_info.add_run(f'{datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n')
    footer_info.add_run('Technology: ').bold = True
    footer_info.add_run('Next.js 16 • React 19 • TypeScript • Tailwind CSS 4\n')
    footer_info.add_run('Repository: ').bold = True
    footer_info.add_run('github.com/Svpriyaa2808/motorcycle')

    # Save the document
    filename = "Motorcycle_Repair_Shops_Project_Documentation.docx"
    doc.save(filename)
    print(f"✓ Word documentation generated successfully: {filename}")
    return filename

if __name__ == "__main__":
    create_word_documentation()
